from pathlib import Path
import math
import shutil
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_IMAGES = ROOT.parent / "Instruction_Images"
DEPTH_SCREENSHOT = Path(r"C:\Users\Raymond Arellano\Pictures\Screenshots\Screenshot 2026-05-26 012606.png")
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "user-guide-assets"
DOCX_PATH = OUT_DIR / "Neverwinter_Forge_User_Guide_v1_1.docx"
PDF_PATH = OUT_DIR / "Neverwinter_Forge_User_Guide_v1_1.pdf"
PAGE_DIR = OUT_DIR / "rendered-user-guide-pages"

GOLD = RGBColor(214, 171, 73)
DARK_GOLD = RGBColor(119, 86, 24)
INK = RGBColor(24, 27, 31)
MUTED = RGBColor(93, 99, 108)
PANEL = "F3F0E8"
PALE_GOLD = "FBF2D7"
LIGHT_GRAY = "F5F6F7"


def font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf") if bold else Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_numbered_box(draw, box, number, image_size):
    w, h = image_size
    scale = max(w, h) / 2200
    line = max(5, int(7 * scale))
    radius = max(34, int(42 * scale))
    x1, y1, x2, y2 = box
    gold = (232, 188, 86, 255)
    shadow = (0, 0, 0, 160)
    draw.rounded_rectangle((x1 + line, y1 + line, x2 + line, y2 + line), radius=18, outline=shadow, width=line)
    draw.rounded_rectangle(box, radius=18, outline=gold, width=line)
    cx = min(max(x1, radius + 8), w - radius - 8)
    cy = min(max(y1, radius + 8), h - radius - 8)
    draw.ellipse((cx - radius, cy - radius, cx + radius, cy + radius), fill=(12, 13, 15, 235), outline=gold, width=line)
    text = str(number)
    fnt = font(max(34, int(42 * scale)), bold=True)
    bbox = draw.textbbox((0, 0), text, font=fnt)
    draw.text((cx - (bbox[2] - bbox[0]) / 2, cy - (bbox[3] - bbox[1]) / 2 - 3), text, fill=(255, 241, 190), font=fnt)


def annotate_image(src_name, out_name, boxes):
    src = Path(src_name)
    if not src.is_absolute():
        src = SOURCE_IMAGES / src_name
    out = ASSET_DIR / out_name
    im = Image.open(src).convert("RGBA")
    overlay = Image.new("RGBA", im.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    for index, box in enumerate(boxes, start=1):
        draw_numbered_box(draw, box, index, im.size)
    Image.alpha_composite(im, overlay).convert("RGB").save(out, quality=94)
    return out


def crop_center(src_name, out_name, crop_box):
    src = SOURCE_IMAGES / src_name
    out = ASSET_DIR / out_name
    im = Image.open(src).convert("RGB")
    im.crop(crop_box).save(out, quality=94)
    return out


def fit_image(path, max_width=None, max_height=None):
    im = Image.open(path)
    w, h = im.size
    if max_width is None:
        max_width = 6.7
    width = max_width
    height = width * h / w
    if max_height and height > max_height:
        height = max_height
        width = height * w / h
    return width, height


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D6AB49", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color="D6AB49", size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_run(paragraph, text, size=10.5, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    add_run(p, text, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    if level == 1:
        add_run(p, text.upper(), size=16, color=DARK_GOLD, bold=True)
        paragraph_border_bottom(p, color="D6AB49", size="10")
    elif level == 2:
        add_run(p, text, size=13, color=INK, bold=True)
    else:
        add_run(p, text, size=11.5, color=DARK_GOLD, bold=True)
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_border(cell, "D6AB49", "8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title + ": ", size=10.5, color=DARK_GOLD, bold=True)
    add_run(p, body, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        p.runs.clear()
        add_run(p, item, size=10.2, color=INK)


def add_numbered_legend(doc, items):
    for number, text in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, f"{number}. ", size=10.2, color=DARK_GOLD, bold=True)
        add_run(p, text, size=10.2, color=INK)


def add_image(doc, path, max_width=6.7, max_height=None, caption=None):
    width, height = fit_image(path, max_width=max_width, max_height=max_height)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        add_run(cp, caption, size=8.8, color=MUTED, italic=True)


def add_two_image_row(doc, left_path, right_path, left_caption, right_caption):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col in table.columns:
        col.width = Inches(3.25)
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "FFFFFF")
            set_cell_border(cell, "E0D6BF", "4")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for idx, path in enumerate((left_path, right_path)):
        p = table.cell(0, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        w, h = fit_image(path, max_width=3.05, max_height=3.4)
        p.add_run().add_picture(str(path), width=Inches(w), height=Inches(h))
    for idx, caption in enumerate((left_caption, right_caption)):
        p = table.cell(1, idx).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, caption, size=9.2, color=MUTED, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_page_break(doc):
    doc.add_page_break()


def prepare_assets():
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    logo_source = SOURCE_IMAGES / "NeverwinterForge_Main_Header.png"
    if not logo_source.exists():
        logo_source = ROOT.parent / "NeverwinterForge_Main_Header.png"
    shutil.copy2(logo_source, ASSET_DIR / "NeverwinterForge_Main_Header.png")

    basic = annotate_image(
        "Using_Preset+Misc_Option.png",
        "annotated_basic_generation.jpg",
        [
            (2670, 90, 3185, 170),
            (3250, 38, 3785, 565),
            (38, 258, 1595, 1485),
            (1625, 258, 3180, 1485),
            (1648, 1505, 3160, 1577),
            (1648, 1590, 3155, 1684),
        ],
    )
    open_prompt = annotate_image(
        "Using_Open_Prompt.png",
        "annotated_open_prompt.jpg",
        [
            (68, 1738, 1295, 1810),
            (92, 775, 1548, 1260),
            (1655, 102, 3165, 1270),
            (3260, 1220, 3785, 1745),
            (3260, 1868, 3785, 1945),
        ],
    )
    upscale = annotate_image(
        "Using_Upscale.png",
        "annotated_upscale.jpg",
        [
            (45, 700, 1705, 775),
            (45, 790, 1705, 865),
            (42, 920, 1935, 1788),
            (43, 1805, 264, 1870),
            (1995, 524, 2445, 804),
        ],
    )
    bas_text = annotate_image(
        "Bas_Relief_Text_to_Image.png",
        "annotated_bas_relief_text.jpg",
        [
            (40, 350, 1575, 940),
            (2030, 362, 2775, 1485),
            (3255, 415, 3782, 1815),
            (1645, 1538, 3150, 1640),
        ],
    )
    depth_maps = annotate_image(
        str(DEPTH_SCREENSHOT) if DEPTH_SCREENSHOT.exists() else "Using_Upscale.png",
        "annotated_depth_normal_maps.jpg",
        [
            (2110, 54, 2440, 145),
            (60, 165, 2435, 910),
            (57, 922, 2442, 1120),
            (57, 1135, 2442, 1228),
            (57, 1290, 1235, 1900),
            (1262, 1290, 2428, 1900),
            (57, 1848, 575, 1935),
        ],
    )
    return {
        "logo": ASSET_DIR / "NeverwinterForge_Main_Header.png",
        "basic": basic,
        "open_prompt": open_prompt,
        "upscale": upscale,
        "bas_text": bas_text,
        "depth_maps": depth_maps,
    }


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(hp, "Neverwinter Forge User Guide", size=8.5, color=MUTED)
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(fp, "Generated images are automatically saved to the app's local outputs folder.", size=8, color=MUTED)
    return doc


def make_docx():
    assets = prepare_assets()
    doc = setup_document()

    add_image(doc, assets["logo"], max_width=6.5, max_height=1.1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "User Guide and Quick Start", size=24, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Version 1.1 - image generation, preset workflows, upscaling, and local output handling", size=11, color=MUTED)
    add_note(
        doc,
        "Current testing recommendation",
        "OpenAI has produced the most consistent front/back matching and prompt adherence in the current Neverwinter Forge tests. Gemini remains available for users who prefer it or want to experiment with Gemini-specific image models.",
    )
    add_heading(doc, "What Neverwinter Forge Does", 1)
    add_para(
        doc,
        "Neverwinter Forge is a focused image-generation front end for creating game-art references, sculpting references, modular outfits, object concepts, bas-relief plates, and post-production upscales without asking the user to manually assemble prompts every time.",
    )
    add_bullets(
        doc,
        [
            "Choose a preset workflow, load an input image or type a concept, then generate through OpenAI, Gemini, or Mock mode.",
            "Use view buttons such as Back View or 3/4 View after a front result is created.",
            "Run post-production tools such as upscaling after you are happy with an image.",
            "Generated images auto-save to the local outputs folder beside the running app or EXE.",
        ],
    )
    add_note(
        doc,
        "Work in progress",
        "The local depth and normal map workflow depends on ComfyUI Desktop and the required models/nodes. It is intended as an advanced local workflow and may continue to evolve.",
    )

    add_page_break(doc)
    add_heading(doc, "Quick Start", 1)
    add_bullets(
        doc,
        [
            "Open Neverwinter Forge and wait for the splash screen to finish.",
            "Choose Mock, Gemini, or OpenAI in the provider selector.",
            "Paste the matching API key if using Gemini or OpenAI.",
            "Choose a preset from Preset Prompts, or use Open Prompt for a direct custom generation.",
            "Load an input image when the workflow is image-to-image, or type the subject when the workflow is text-to-image.",
            "Choose output size, aspect ratio, and quality options when available.",
            "Click Generate. The output appears in the preview and is automatically saved to the outputs folder.",
        ],
    )
    add_heading(doc, "Provider Guidance", 2)
    add_bullets(
        doc,
        [
            "OpenAI is recommended for the strongest prompt adherence in the current test set, especially front/back matching for outfits and controlled view generation.",
            "Gemini can still produce excellent results, especially for supported bas-relief and concept workflows, but matching a second view to a first view can vary by model and prompt.",
            "Mock mode is free and is meant for testing the interface. It does not spend API credits and should not be judged as real image generation.",
            "Cost estimates are approximate. Provider pricing can change, and image or text input costs may vary.",
        ],
    )
    add_heading(doc, "Where Outputs Go", 2)
    add_bullets(
        doc,
        [
            "Every successful generation auto-saves an image file and a matching metadata JSON file.",
            "The Open Outputs Folder button opens the output folder used by the currently running copy of the app.",
            "Save Browser Copy is only a manual browser download option. The autosaved file is the one users should rely on for normal work.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Main Generation Screen", 1)
    add_image(doc, assets["basic"], max_width=6.75, caption="Main image-to-image workflow with preset, input, output, view buttons, and autosave area.")
    add_numbered_legend(
        doc,
        [
            "Provider selector: choose Mock, Gemini, or OpenAI before generating.",
            "Preset Prompts: choose the workflow template. The active preset controls the main prompt.",
            "Input panel: drop or choose the source image for image-to-image workflows.",
            "Output panel: generated result appears here after completion.",
            "View buttons: create follow-up views from the current front result when available.",
            "Autosave panel: confirms the image and metadata file saved to the outputs folder.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Preset + Misc", 1)
    add_para(
        doc,
        "Preset + Misc is the everyday mode for most work. The preset supplies the structured production prompt, while the Misc / Extra Prompt box adds a small extra instruction on top.",
    )
    add_bullets(
        doc,
        [
            "Use Preset + Misc when you want the normal workflow plus a small extra direction.",
            "Use Preset Only when you want the template to run cleanly with no extra guidance.",
            "Use Misc Override Only when you intentionally want your own text to drive the result instead of the preset.",
        ],
    )
    add_note(
        doc,
        "Practical tip",
        "For production consistency, keep Misc short. One or two focused additions usually behave better than a second full prompt.",
    )

    add_heading(doc, "D&D Miniature", 2)
    add_para(
        doc,
        "This workflow converts a character reference into a monochrome orthographic miniature statue reference. It is designed for downstream image-to-3D and sculpting workflows.",
    )
    add_bullets(
        doc,
        [
            "Use a clear full-body input when possible.",
            "Generate the front view first.",
            "Use Back View or 3/4 View only after the front result is acceptable.",
            "OpenAI is currently preferred for the most stable follow-up views.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Modular Outfit Workflows", 1)
    add_para(
        doc,
        "The modular outfit presets convert character outfit references into clean T-pose style outfit renders suitable for modeling reference. Separate male and female presets use different locked silhouette references.",
    )
    add_two_image_row(
        doc,
        SOURCE_IMAGES / "Modular_Outfit_Male_Front.png",
        SOURCE_IMAGES / "Modular_Outfit_Male_Back.png",
        "Modular Outfit - Male, front view",
        "Modular Outfit - Male, back view",
    )
    add_two_image_row(
        doc,
        SOURCE_IMAGES / "Modular_Outfit_Female_Front.png",
        SOURCE_IMAGES / "Modular_Outfit_Female_Back.png",
        "Modular Outfit - Female, front view",
        "Modular Outfit - Female, back view",
    )
    add_bullets(
        doc,
        [
            "Use the male preset for male outfit silhouettes and the female preset for female outfit silhouettes.",
            "Cloaks, identifying symbols, and overly specific insignias should be simplified or removed by the prompt, but review the result before moving downstream.",
            "If a cloak or emblem survives, rerun with a short Misc instruction such as remove cloak and remove heraldic symbols.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Object Concept", 1)
    add_para(
        doc,
        "Object Concept creates clean orthographic object renders from text. It is useful for props that need to move into AI-to-3D, Blender cleanup, game asset creation, or Substance Painter reference work.",
    )
    add_two_image_row(
        doc,
        SOURCE_IMAGES / "Object_Concept_Front.png",
        SOURCE_IMAGES / "Object_Concept_Back.png",
        "Object Concept, front view",
        "Object Concept, back view",
    )
    add_bullets(
        doc,
        [
            "Type the object you want to generate in the object prompt field.",
            "Generate the front view first, then use Back View or Side View when needed.",
            "Keep object descriptions concrete: material, shape, era, and intended use are more helpful than broad style words.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Bas-Relief Emblem and Concept", 1)
    add_image(doc, assets["bas_text"], max_width=6.75, caption="Bas-Relief Concept text-to-image workflow with output size and aspect-ratio controls.")
    add_numbered_legend(
        doc,
        [
            "Text prompt: describe the bas-relief subject or scene.",
            "Output preview: the generated carved plate appears here.",
            "Gemini canvas controls: set model, output size, and bas-relief aspect ratio when using Gemini.",
            "Autosave panel: confirms the generated image and metadata file.",
        ],
    )
    add_bullets(
        doc,
        [
            "Bas-Relief Emblem uses an input image as the visual source.",
            "Bas-Relief Concept uses written text as the visual source.",
            "Use the aspect ratio selector before generating; it tells Gemini to compose the plate in that canvas ratio rather than copying the input image ratio.",
            "Use these outputs as height-map, sculpting, carving, or bas-relief modeling references.",
        ],
    )
    add_two_image_row(
        doc,
        SOURCE_IMAGES / "Bas_Relief_from_Image.png",
        SOURCE_IMAGES / "Bas_Relief_from Text.jpg",
        "Bas-relief generated from an image reference",
        "Bas-relief generated from text",
    )

    add_page_break(doc)
    add_heading(doc, "Open Prompt", 1)
    add_image(doc, assets["open_prompt"], max_width=6.75, caption="Open Prompt is the flexible no-preset mode for custom generations.")
    add_numbered_legend(
        doc,
        [
            "Open Prompt button: activates direct prompt generation.",
            "Prompt field: write the full prompt or instruction.",
            "Output preview: generated result appears here.",
            "Provider settings: choose model, output size, aspect ratio, quality, and review estimated cost.",
            "Generate Open Prompt: starts generation with your direct prompt.",
        ],
    )
    add_note(
        doc,
        "When to use it",
        "Open Prompt is best when the preset system is too specific, when testing a new idea, or when demonstrating the app as a general image-generation tool.",
    )

    add_page_break(doc)
    add_heading(doc, "Upscaler", 1)
    add_image(doc, assets["upscale"], max_width=6.75, caption="Upscaler workflow for post-production 2x or 4x output.")
    add_numbered_legend(
        doc,
        [
            "Use Current Output: loads the most recent generated image into the upscaler.",
            "Upscale 2x / 4x: starts local upscaling after the image is loaded.",
            "Comparison area: shows original and upscaled images side by side.",
            "Save Upscaled: saves or downloads the upscaled version for manual use.",
            "Local status: shows installation and processing messages.",
        ],
    )
    add_bullets(
        doc,
        [
            "Upscaling is a post-production step. Generate the image first, then upscale only the version you want to keep.",
            "Use 2x for a fast quality boost and 4x when you need a larger reference image.",
            "Upscaled files are also written into the local outputs folder.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Depth and Normal Maps", 1)
    add_image(doc, assets["depth_maps"], max_width=6.75, caption="Depth / Normal Maps workflow with source image, dependency status, generated depth map, and generated normal map.")
    add_numbered_legend(
        doc,
        [
            "Refresh: checks whether the local ComfyUI workflow and dependencies are ready.",
            "Input preview: use the current output or choose an image for map generation.",
            "Model dependencies: show or hide local model/node status.",
            "Generate Depth / Normal Maps: sends the image to the local ComfyUI workflow.",
            "Depth result: grayscale depth/height-style output.",
            "Normal result: RGB normal map output.",
            "Save buttons: manually save depth or normal maps when needed.",
        ],
    )
    add_para(
        doc,
        "Depth / Normal Maps is a local workflow that uses ComfyUI Desktop, the required custom nodes, and the expected Lotus/VAE model files. It is not an API generation mode.",
    )
    add_bullets(
        doc,
        [
            "Install ComfyUI Desktop first if you plan to use local depth or normal map generation.",
            "Run the included install_depth_models.bat helper to place the expected models and nodes into the detected ComfyUI folders.",
            "Start ComfyUI Desktop before running the depth/normal workflow.",
            "Use the Refresh or dependency panel to check whether Neverwinter Forge sees the local setup.",
        ],
    )
    add_note(
        doc,
        "Advanced feature",
        "This local workflow is more technical than the API presets. It is included for users who want mesh-support maps, but the main image workflows do not require ComfyUI.",
    )

    add_heading(doc, "Troubleshooting", 1)
    add_bullets(
        doc,
        [
            "No API output: confirm the selected provider has the matching API key filled in.",
            "Wrong style: choose Preset Only to remove any accidental Misc instructions, then test again.",
            "Aspect ratio ignored: make sure the bas-relief or Open Prompt aspect ratio was selected before pressing Generate.",
            "Second view does not match: regenerate from the best front result and prefer OpenAI for view-matching work.",
            "No output folder visible: click Open Outputs Folder. The active app creates and uses the outputs folder beside the running EXE.",
            "ComfyUI offline: launch ComfyUI Desktop first, then refresh the Depth / Normal Maps status.",
        ],
    )

    add_heading(doc, "Release Notes for Users", 1)
    add_para(
        doc,
        "Neverwinter Forge is actively evolving. Presets, provider behavior, local model workflows, and UI polish may continue to improve as new workflows are tested.",
    )
    add_para(
        doc,
        "For best results, keep a small set of successful outputs as references, use the autosaved files in outputs for downstream work, and rerun with short targeted instructions when a generated image preserves an unwanted cloak, symbol, or design detail.",
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


PAGE_W, PAGE_H = 1700, 2200
MARGIN = 100
BG = (9, 10, 11)
PANEL_BG = (28, 31, 36)
PANEL_EDGE = (199, 152, 55)
TEXT = (235, 238, 242)
TEXT_MUTED = (172, 180, 190)
GOLD_RGB = (224, 177, 74)
GOLD_SOFT = (246, 219, 147)


def draw_wrapped(draw, text, xy, fnt, fill=TEXT, max_width=900, line_gap=8):
    x, y = xy
    words = text.split()
    lines = []
    line = ""
    for word in words:
        trial = (line + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width or not line:
            line = trial
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    line_h = draw.textbbox((0, 0), "Ag", font=fnt)[3] + line_gap
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h
    return y


def new_pdf_page(title=None):
    page = Image.new("RGB", (PAGE_W, PAGE_H), BG)
    draw = ImageDraw.Draw(page)
    draw.rectangle((0, 0, PAGE_W, 18), fill=(40, 30, 12))
    if title:
        draw.text((MARGIN, 58), title.upper(), font=font(48, bold=True), fill=GOLD_RGB)
        draw.line((MARGIN, 126, PAGE_W - MARGIN, 126), fill=PANEL_EDGE, width=3)
    return page, draw


def image_fit_size(path, max_w, max_h):
    im = Image.open(path).convert("RGB")
    ratio = min(max_w / im.width, max_h / im.height)
    return im.resize((int(im.width * ratio), int(im.height * ratio)), Image.LANCZOS)


def paste_image(draw, page, path, box, caption=None):
    x1, y1, x2, y2 = box
    im = image_fit_size(path, x2 - x1, y2 - y1)
    x = x1 + (x2 - x1 - im.width) // 2
    y = y1 + (y2 - y1 - im.height) // 2
    draw.rounded_rectangle((x - 10, y - 10, x + im.width + 10, y + im.height + 10), radius=16, outline=PANEL_EDGE, width=3, fill=(14, 15, 17))
    page.paste(im, (x, y))
    if caption:
        draw.text((x1, y2 + 18), caption, font=font(26), fill=TEXT_MUTED)


def panel(draw, xyxy, title=None):
    draw.rounded_rectangle(xyxy, radius=24, fill=PANEL_BG, outline=PANEL_EDGE, width=3)
    if title:
        draw.text((xyxy[0] + 32, xyxy[1] + 28), title.upper(), font=font(32, bold=True), fill=GOLD_SOFT)


def bullet_list(draw, items, x, y, max_width, fnt=None):
    if fnt is None:
        fnt = font(29)
    bullet_fnt = font(30, bold=True)
    for item in items:
        draw.text((x, y + 3), "-", font=bullet_fnt, fill=GOLD_RGB)
        y = draw_wrapped(draw, item, (x + 34, y), fnt, fill=TEXT, max_width=max_width - 34, line_gap=8) + 14
    return y


def numbered_list(draw, items, x, y, max_width):
    for i, item in enumerate(items, start=1):
        r = 22
        draw.ellipse((x, y + 4, x + r * 2, y + 4 + r * 2), fill=(13, 14, 15), outline=GOLD_RGB, width=3)
        label = str(i)
        f = font(26, bold=True)
        bb = draw.textbbox((0, 0), label, font=f)
        draw.text((x + r - (bb[2] - bb[0]) / 2, y + 4 + r - (bb[3] - bb[1]) / 2 - 2), label, font=f, fill=GOLD_SOFT)
        y = draw_wrapped(draw, item, (x + 62, y), font(28), fill=TEXT, max_width=max_width - 62, line_gap=8) + 18
    return y


def footer(draw, page_no):
    draw.line((MARGIN, PAGE_H - 86, PAGE_W - MARGIN, PAGE_H - 86), fill=(69, 56, 30), width=2)
    draw.text((MARGIN, PAGE_H - 62), "Neverwinter Forge User Guide", font=font(22), fill=TEXT_MUTED)
    draw.text((PAGE_W - MARGIN - 90, PAGE_H - 62), f"{page_no:02d}", font=font(22), fill=TEXT_MUTED)


def make_pdf():
    assets = prepare_assets()
    PAGE_DIR.mkdir(exist_ok=True)
    for old in PAGE_DIR.glob("page-*.png"):
        old.unlink()
    pages = []

    page, draw = new_pdf_page()
    logo = image_fit_size(assets["logo"], 1420, 220)
    page.paste(logo, ((PAGE_W - logo.width) // 2, 180))
    draw.text((MARGIN, 520), "USER GUIDE", font=font(82, bold=True), fill=TEXT)
    draw.text((MARGIN, 620), "Version 1.1", font=font(42, bold=True), fill=GOLD_RGB)
    panel(draw, (MARGIN, 760, PAGE_W - MARGIN, 1115), "What this app is for")
    draw_wrapped(
        draw,
        "Neverwinter Forge is a focused image-generation front end for creating game-art references, miniature statue concepts, modular outfit sheets, object concepts, bas-relief plates, and post-production upscales without manually rebuilding long prompts.",
        (MARGIN + 34, 830),
        font(34),
        max_width=PAGE_W - 2 * MARGIN - 68,
        line_gap=10,
    )
    panel(draw, (MARGIN, 1210, PAGE_W - MARGIN, 1575), "Current testing recommendation")
    draw_wrapped(
        draw,
        "OpenAI has produced the most consistent front/back matching and prompt adherence in current testing. Gemini remains available and can produce strong results, especially when users want to experiment with Gemini-specific image models.",
        (MARGIN + 34, 1280),
        font(31),
        max_width=PAGE_W - 2 * MARGIN - 68,
        line_gap=10,
    )
    panel(draw, (MARGIN, 1670, PAGE_W - MARGIN, 1930), "Developer note")
    draw_wrapped(
        draw,
        "Some local workflows, especially depth and normal maps through ComfyUI Desktop, are advanced features and may continue to evolve.",
        (MARGIN + 34, 1744),
        font(31),
        max_width=PAGE_W - 2 * MARGIN - 68,
        line_gap=10,
    )
    footer(draw, 1)
    pages.append(page)

    page, draw = new_pdf_page("Quick Start")
    panel(draw, (MARGIN, 180, PAGE_W - MARGIN, 885), "Start here")
    bullet_list(
        draw,
        [
            "Open Neverwinter Forge and wait for the splash screen to finish.",
            "Choose Mock, Gemini, or OpenAI. Paste the matching API key for Gemini or OpenAI.",
            "Choose a preset, or use Open Prompt for a direct custom generation.",
            "Load an input image for image-to-image workflows, or type a subject for text-to-image workflows.",
            "Set size, aspect ratio, and quality controls when they are available.",
            "Click Generate. The result appears in Output and auto-saves beside the running app.",
        ],
        MARGIN + 36,
        255,
        PAGE_W - 2 * MARGIN - 72,
    )
    panel(draw, (MARGIN, 960, PAGE_W - MARGIN, 1360), "Provider guidance")
    bullet_list(
        draw,
        [
            "OpenAI is recommended when view matching and strict prompt adherence matter most.",
            "Gemini is optional and useful for supported model experiments, but second-view matching can vary.",
            "Mock mode is for interface testing only and does not spend API credits.",
        ],
        MARGIN + 36,
        1035,
        PAGE_W - 2 * MARGIN - 72,
    )
    panel(draw, (MARGIN, 1435, PAGE_W - MARGIN, 1840), "Output folder")
    bullet_list(
        draw,
        [
            "Generated images and matching metadata JSON files are automatically saved.",
            "Use Open Outputs Folder to open the folder used by the currently running EXE.",
            "Save Browser Copy is a manual browser download option, not the primary autosave path.",
        ],
        MARGIN + 36,
        1510,
        PAGE_W - 2 * MARGIN - 72,
    )
    footer(draw, 2)
    pages.append(page)

    page, draw = new_pdf_page("Main Generation Screen")
    paste_image(draw, page, assets["basic"], (MARGIN, 175, PAGE_W - MARGIN, 1290), "Preset image-to-image workflow with numbered click targets.")
    numbered_list(
        draw,
        [
            "Provider selector: choose Mock, Gemini, or OpenAI.",
            "Preset Prompts: choose the workflow template.",
            "Input panel: load the source image.",
            "Output panel: review the generated result.",
            "View buttons: create follow-up back or angled views when available.",
            "Autosave panel: confirms the saved output file and metadata.",
        ],
        MARGIN,
        1360,
        PAGE_W - 2 * MARGIN,
    )
    footer(draw, 3)
    pages.append(page)

    page, draw = new_pdf_page("Preset + Misc")
    panel(draw, (MARGIN, 180, PAGE_W - MARGIN, 640), "Prompt modes")
    bullet_list(
        draw,
        [
            "Preset + Misc: use the preset prompt and add a small extra instruction.",
            "Preset Only: run the preset cleanly with no extra direction.",
            "Misc Override Only: intentionally let your custom prompt drive the generation.",
        ],
        MARGIN + 36,
        260,
        PAGE_W - 2 * MARGIN - 72,
    )
    panel(draw, (MARGIN, 720, PAGE_W - MARGIN, 1260), "D&D Miniature")
    bullet_list(
        draw,
        [
            "Use a clear full-body character reference when possible.",
            "Generate the front miniature first.",
            "Use Back View or 3/4 View after the front result is acceptable.",
            "OpenAI is currently preferred for consistent follow-up views.",
        ],
        MARGIN + 36,
        800,
        PAGE_W - 2 * MARGIN - 72,
    )
    paste_image(draw, page, SOURCE_IMAGES / "Using_Preset+Misc_Option.png", (MARGIN, 1345, PAGE_W - MARGIN, 1955), "Example: input reference converted into a miniature statue output.")
    footer(draw, 4)
    pages.append(page)

    page, draw = new_pdf_page("Modular Outfit")
    draw_wrapped(
        draw,
        "The modular outfit presets convert outfit references into clean T-pose style outfit renders. Male and female presets use different locked silhouette references.",
        (MARGIN, 160),
        font(31),
        max_width=PAGE_W - 2 * MARGIN,
    )
    paste_image(draw, page, SOURCE_IMAGES / "Modular_Outfit_Male_Front.png", (MARGIN, 310, 835, 1035), "Male front")
    paste_image(draw, page, SOURCE_IMAGES / "Modular_Outfit_Male_Back.png", (865, 310, PAGE_W - MARGIN, 1035), "Male back")
    paste_image(draw, page, SOURCE_IMAGES / "Modular_Outfit_Female_Front.png", (MARGIN, 1150, 835, 1875), "Female front")
    paste_image(draw, page, SOURCE_IMAGES / "Modular_Outfit_Female_Back.png", (865, 1150, PAGE_W - MARGIN, 1875), "Female back")
    bullet_list(
        draw,
        [
            "Review generated outfits for unwanted cloaks, heraldic symbols, or identifiers.",
            "If an unwanted item remains, rerun with a short Misc instruction to remove it.",
        ],
        MARGIN,
        1970,
        PAGE_W - 2 * MARGIN,
        fnt=font(25),
    )
    footer(draw, 5)
    pages.append(page)

    page, draw = new_pdf_page("Object Concept")
    draw_wrapped(
        draw,
        "Object Concept is a text-to-image workflow for clean orthographic prop references. It is useful for AI-to-3D, Blender cleanup, game asset creation, and Substance Painter reference work.",
        (MARGIN, 160),
        font(31),
        max_width=PAGE_W - 2 * MARGIN,
    )
    paste_image(draw, page, SOURCE_IMAGES / "Object_Concept_Front.png", (MARGIN, 350, 835, 1185), "Object front")
    paste_image(draw, page, SOURCE_IMAGES / "Object_Concept_Back.png", (865, 350, PAGE_W - MARGIN, 1185), "Object back")
    panel(draw, (MARGIN, 1305, PAGE_W - MARGIN, 1775), "How to write the object")
    bullet_list(
        draw,
        [
            "Name the object plainly.",
            "Add material, era, shape language, and intended use.",
            "Generate the front first, then use back or side view when needed.",
        ],
        MARGIN + 36,
        1380,
        PAGE_W - 2 * MARGIN - 72,
    )
    footer(draw, 6)
    pages.append(page)

    page, draw = new_pdf_page("Bas-Relief")
    paste_image(draw, page, assets["bas_text"], (MARGIN, 175, PAGE_W - MARGIN, 1110), "Bas-Relief Concept text-to-image workflow.")
    numbered_list(
        draw,
        [
            "Text prompt: describe the bas-relief subject or scene.",
            "Output preview: generated carved plate.",
            "Canvas controls: choose model, output size, and aspect ratio.",
            "Autosave panel: confirms image and metadata saved to outputs.",
        ],
        MARGIN,
        1165,
        PAGE_W - 2 * MARGIN,
    )
    paste_image(draw, page, SOURCE_IMAGES / "Bas_Relief_from_Image.png", (MARGIN, 1570, 835, 2010), "From image")
    paste_image(draw, page, SOURCE_IMAGES / "Bas_Relief_from Text.jpg", (865, 1570, PAGE_W - MARGIN, 2010), "From text")
    footer(draw, 7)
    pages.append(page)

    page, draw = new_pdf_page("Open Prompt")
    paste_image(draw, page, assets["open_prompt"], (MARGIN, 175, PAGE_W - MARGIN, 1235), "Open Prompt is the flexible no-preset mode.")
    numbered_list(
        draw,
        [
            "Open Prompt button activates direct prompt generation.",
            "Prompt field holds your custom instruction.",
            "Output preview shows the result.",
            "Provider settings control model, size, aspect ratio, quality, and estimated cost.",
            "Generate Open Prompt starts the job.",
        ],
        MARGIN,
        1305,
        PAGE_W - 2 * MARGIN,
    )
    footer(draw, 8)
    pages.append(page)

    page, draw = new_pdf_page("Upscaler")
    paste_image(draw, page, assets["upscale"], (MARGIN, 175, PAGE_W - MARGIN, 1320), "Local 2x and 4x post-production upscaling.")
    numbered_list(
        draw,
        [
            "Use Current Output loads the last generated image.",
            "Upscale 2x or 4x starts local upscaling.",
            "Comparison area shows original and upscaled versions.",
            "Save Upscaled lets the user save the upscaled image manually.",
            "Local status confirms install and processing messages.",
        ],
        MARGIN,
        1390,
        PAGE_W - 2 * MARGIN,
    )
    footer(draw, 9)
    pages.append(page)

    page, draw = new_pdf_page("Depth / Normal Maps")
    paste_image(draw, page, assets["depth_maps"], (MARGIN, 175, PAGE_W - MARGIN, 1135), "Local ComfyUI depth and normal map workflow.")
    numbered_list(
        draw,
        [
            "Refresh checks the local ComfyUI workflow status.",
            "Input preview shows the image being processed.",
            "Model dependencies can be expanded when troubleshooting.",
            "Generate Depth / Normal Maps starts the local workflow.",
            "Depth result is the grayscale height/depth-style output.",
            "Normal result is the RGB normal map output.",
            "Save buttons manually save the generated maps.",
        ],
        MARGIN,
        1205,
        PAGE_W - 2 * MARGIN,
    )
    footer(draw, 10)
    pages.append(page)

    page, draw = new_pdf_page("Troubleshooting")
    panel(draw, (MARGIN, 180, PAGE_W - MARGIN, 945), "Depth / Normal Maps")
    bullet_list(
        draw,
        [
            "This is a local workflow that depends on ComfyUI Desktop.",
            "Run install_depth_models.bat to install or check the expected model and node files.",
            "Start ComfyUI Desktop before running depth or normal maps.",
            "Use Refresh or the dependency panel to check what Neverwinter Forge can see.",
        ],
        MARGIN + 36,
        255,
        PAGE_W - 2 * MARGIN - 72,
    )
    panel(draw, (MARGIN, 1035, PAGE_W - MARGIN, 1845), "Common fixes")
    bullet_list(
        draw,
        [
            "No API output: confirm the selected provider has the matching API key.",
            "Wrong style: try Preset Only to remove accidental Misc instructions.",
            "Aspect ratio ignored: choose aspect ratio before pressing Generate.",
            "Second view does not match: regenerate from the best front result and prefer OpenAI for view matching.",
            "No output folder visible: click Open Outputs Folder. The active app creates outputs beside the running EXE.",
            "ComfyUI offline: launch ComfyUI Desktop first, then refresh the local tool status.",
        ],
        MARGIN + 36,
        1110,
        PAGE_W - 2 * MARGIN - 72,
        fnt=font(27),
    )
    footer(draw, 11)
    pages.append(page)

    page_paths = []
    for idx, page in enumerate(pages, start=1):
        path = PAGE_DIR / f"page-{idx:02d}.png"
        page.save(path)
        page_paths.append(path)

    rgb_pages = [Image.open(path).convert("RGB") for path in page_paths]
    rgb_pages[0].save(PDF_PATH, save_all=True, append_images=rgb_pages[1:], resolution=200.0)
    print(PDF_PATH)
    return PDF_PATH, page_paths


if __name__ == "__main__":
    print(make_docx())
    make_pdf()
